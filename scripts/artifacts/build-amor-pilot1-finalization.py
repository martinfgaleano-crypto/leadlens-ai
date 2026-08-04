#!/usr/bin/env python3
# Premium Amor de Gea Pilot 1 customer deliverables (final release build).
# Renders 4 customer-safe files from output/amor-pilot1-deliverable.data.json
# (single source of truth exported from the approved intelligence modules).
# Compressed page economy, coherent numbering (cover excluded, "Página X de N"
# body count), drawn empty checkboxes, evaluation guides in the feedback doc.
# Deterministic output (stable download checksums). No internal version tokens,
# no named buyers, no timing claims. Run: python build-amor-pilot1-finalization.py
import json
import shutil
from datetime import datetime
from pathlib import Path

import reportlab.rl_config
reportlab.rl_config.invariant = 1  # deterministic PDFs so download checksums stay stable
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, PageBreak, Table,
    TableStyle, KeepTogether, Flowable, ListFlowable, ListItem, NextPageTemplate,
)
from reportlab.pdfgen import canvas as canvasmod
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
DATA = json.loads((ROOT / "output/amor-pilot1-deliverable.data.json").read_text(encoding="utf-8"))
PDF = ROOT / "output/pdf"
DOCX = ROOT / "output/docx"
PUBLIC = ROOT / "public/pilot-deliverables"
for d in (PDF, DOCX, PUBLIC):
    d.mkdir(parents=True, exist_ok=True)

NAVY = colors.HexColor("#17352C")
GREEN = colors.HexColor("#4E6A54")
GOLD = colors.HexColor("#B48A4A")
GOLD_SOFT = colors.HexColor("#D8C29A")
SAGE = colors.HexColor("#EAF0E9")
SAGE_LINE = colors.HexColor("#CAD8CD")
INK = colors.HexColor("#24332C")
MUTED = colors.HexColor("#6B7873")
CREAM = colors.HexColor("#F8F5ED")
CARD = colors.HexColor("#FBFAF5")
PAGE_W, PAGE_H = letter
LM = RM = 56
CONTENT_W = PAGE_W - LM - RM

styles = getSampleStyleSheet()
def _add(name, **kw):
    styles.add(ParagraphStyle(name=name, **kw))
_add("Kicker", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=GOLD, spaceAfter=3)
_add("H1x", fontName="Helvetica-Bold", fontSize=19, leading=23, textColor=NAVY, spaceAfter=4)
_add("H2x", fontName="Helvetica-Bold", fontSize=12, leading=15, textColor=NAVY, spaceBefore=9, spaceAfter=3)
_add("H3x", fontName="Helvetica-Bold", fontSize=9.6, leading=12.5, textColor=GREEN, spaceBefore=6, spaceAfter=2)
_add("Body", fontName="Helvetica", fontSize=9.6, leading=14, textColor=INK, spaceAfter=6, alignment=TA_JUSTIFY)
_add("BodyL", fontName="Helvetica", fontSize=9.6, leading=14, textColor=INK, spaceAfter=6)
_add("Small", fontName="Helvetica", fontSize=8, leading=11, textColor=MUTED, spaceAfter=3)
_add("CardH", fontName="Helvetica-Bold", fontSize=10.5, leading=13, textColor=NAVY, spaceAfter=2)
_add("CardBody", fontName="Helvetica", fontSize=9, leading=12.5, textColor=INK, spaceAfter=2)
_add("Tag", fontName="Helvetica-Bold", fontSize=7.2, leading=9, textColor=colors.white)
_add("EvLabel", fontName="Helvetica-Bold", fontSize=7, leading=8.5, textColor=GOLD)
_add("EvBody", fontName="Helvetica", fontSize=8.5, leading=11.5, textColor=INK)
_add("TH", fontName="Helvetica-Bold", fontSize=8.5, leading=11, textColor=colors.white)
_add("TD", fontName="Helvetica", fontSize=8.6, leading=11.8, textColor=INK)
_add("Callout", fontName="Helvetica-Bold", fontSize=10.5, leading=15, textColor=NAVY)

def P(text, style="Body"):
    return Paragraph(text, styles[style])

class Rule(Flowable):
    def __init__(self, width=CONTENT_W, color=GOLD, thick=1.4, gap=6):
        super().__init__(); self.width = width; self.color = color; self.thick = thick; self.gap = gap
    def wrap(self, aw, ah):
        return (self.width, self.gap + self.thick)
    def draw(self):
        self.canv.setStrokeColor(self.color); self.canv.setLineWidth(self.thick)
        self.canv.line(0, self.gap, self.width, self.gap)

class Box(Flowable):
    """A genuine empty vector checkbox (never a solid glyph)."""
    def __init__(self, s=9):
        super().__init__(); self.s = s
    def wrap(self, aw, ah):
        return (self.s, self.s)
    def draw(self):
        self.canv.setStrokeColor(INK); self.canv.setLineWidth(0.8)
        self.canv.rect(0, 0, self.s, self.s, fill=0)

def section(story, number, title, kicker=None):
    story.append(P(kicker.upper() if kicker else f"SECCIÓN {number}", "Kicker"))
    story.append(P(title, "H1x"))
    story.append(Rule(width=64, color=GOLD, thick=2.4, gap=5))
    story.append(Spacer(1, 7))

def tag(text, bg):
    t = Table([[P(text, "Tag")]], colWidths=[len(text) * 4.6 + 12])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), bg), ("LEFTPADDING", (0, 0), (-1, -1), 6),
                           ("RIGHTPADDING", (0, 0), (-1, -1), 6), ("TOPPADDING", (0, 0), (-1, -1), 2.5),
                           ("BOTTOMPADDING", (0, 0), (-1, -1), 2.5), ("ROUNDEDCORNERS", [3, 3, 3, 3])]))
    return t

ROUTE_BG = {"Hotelería y spa": colors.HexColor("#3E6B8A"), "Retail especializado": GREEN,
            "Regalos corporativos": GOLD, "Distribución regional": MUTED}

def mk_table(header, rows, widths, zebra=True):
    data = [[P(c, "TH") for c in header]] + [[P(str(c), "TD") for c in r] for r in rows]
    t = Table(data, colWidths=widths, repeatRows=1)
    style = [("BACKGROUND", (0, 0), (-1, 0), NAVY), ("GRID", (0, 0), (-1, -1), 0.4, SAGE_LINE),
             ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 7),
             ("RIGHTPADDING", (0, 0), (-1, -1), 7), ("TOPPADDING", (0, 0), (-1, -1), 5.5),
             ("BOTTOMPADDING", (0, 0), (-1, -1), 5.5)]
    if zebra:
        for i in range(1, len(data)):
            style.append(("BACKGROUND", (0, i), (-1, i), CREAM if i % 2 else CARD))
    t.setStyle(TableStyle(style))
    return t

def table(story, header, rows, widths, zebra=True):
    story.append(mk_table(header, rows, widths, zebra))

def callout(story, text, after=8):
    t = Table([[P(text, "Callout")]], colWidths=[CONTENT_W])
    t.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), SAGE), ("LEFTPADDING", (0, 0), (-1, -1), 13),
                           ("RIGHTPADDING", (0, 0), (-1, -1), 13), ("TOPPADDING", (0, 0), (-1, -1), 10),
                           ("BOTTOMPADDING", (0, 0), (-1, -1), 10), ("LINEBEFORE", (0, 0), (0, -1), 3, GOLD),
                           ("ROUNDEDCORNERS", [4, 4, 4, 4])]))
    story.append(t); story.append(Spacer(1, after))

def bullets(story, items, style="Body"):
    story.append(ListFlowable([ListItem(P(x, style), leftIndent=12, value="•") for x in items],
                              bulletType="bullet", start="•", bulletColor=GOLD, bulletFontSize=8))

def checkbox_row(items, label_style="CardBody", box=9):
    cells, widths, ts = [], [], [("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                 ("LEFTPADDING", (0, 0), (-1, -1), 1), ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                                 ("TOPPADDING", (0, 0), (-1, -1), 1.5), ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5)]
    for label in items:
        cells.append(Box(box)); widths.append(box + 3)
        cells.append(P(str(label), label_style)); widths.append(max(16, len(str(label)) * 4.9 + 10))
    t = Table([cells], colWidths=widths, hAlign="LEFT")
    t.setStyle(TableStyle(ts))
    return t

def evidence_box(ev, w):
    cells = [[P("FUENTE", "EvLabel")], [P(f'{ev["source"]} · sitio oficial', "EvBody")],
             [P(f'HECHO PÚBLICO · consultado {ev["retrieved"]}', "EvLabel")], [P(ev["fact"], "EvBody")],
             [P("CONFIRMA", "EvLabel")], [P(ev["proves"], "EvBody")],
             [P("NO CONFIRMA", "EvLabel")], [P(ev["not_proves"], "EvBody")]]
    et = Table(cells, colWidths=[w])
    et.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.white), ("LEFTPADDING", (0, 0), (-1, -1), 8),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 8), ("TOPPADDING", (0, 0), (-1, -1), 1.5),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 1.5), ("BOX", (0, 0), (-1, -1), 0.5, SAGE_LINE)]))
    return et

def account_card(a, full_evidence=True):
    ev = a["evidence"]
    left = [P(a["name"], "CardH"), tag(a["route"], ROUTE_BG.get(a["route"], GREEN)), Spacer(1, 4),
            P(a["why"], "CardBody"), P(f'<b>Prueba inicial:</b> {a["test"]}', "CardBody"),
            P(f'<b>Por validar:</b> {a["next"]}', "CardBody")]
    if full_evidence:
        row = [[left, evidence_box(ev, CONTENT_W * 0.42 - 16)]]
        widths = [CONTENT_W * 0.58, CONTENT_W * 0.42]
    else:
        left.append(Spacer(1, 2))
        left.append(P(f'<b>Evidencia:</b> {ev["source"]} · {ev["fact"]}', "Small"))
        row = [[left]]; widths = [CONTENT_W]
    card = Table(row, colWidths=widths)
    card.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), CARD), ("VALIGN", (0, 0), (-1, -1), "TOP"),
                              ("LEFTPADDING", (0, 0), (-1, -1), 13), ("RIGHTPADDING", (0, 0), (-1, -1), 13),
                              ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                              ("BOX", (0, 0), (-1, -1), 0.6, SAGE_LINE), ("LINEBEFORE", (0, 0), (0, -1), 3, ROUTE_BG.get(a["route"], GREEN)),
                              ("ROUNDEDCORNERS", [5, 5, 5, 5])]))
    return KeepTogether([card, Spacer(1, 8)])

class PortfolioMap(Flowable):
    def __init__(self, data):
        super().__init__(); self.data = data; self.width = CONTENT_W; self.height = 210
    def wrap(self, aw, ah):
        return (self.width, self.height)
    def draw(self):
        c = self.canv
        cols = [("PRIMERA SECUENCIA", self.data["portfolio"]["first_validation"], NAVY),
                ("PRIORIDAD ESTRATÉGICA", self.data["portfolio"]["strategic_priority"], GREEN),
                ("INVESTIGACIÓN SELECTIVA", self.data["portfolio"]["investigate_selectively"], MUTED)]
        route_of = {a["name"]: a["route"] for a in self.data["accounts"]}
        gap = 12; cw = (self.width - 2 * gap) / 3; top = self.height - 4
        for i, (label, names, col) in enumerate(cols):
            x = i * (cw + gap)
            c.setFillColor(col); c.roundRect(x, top - 21, cw, 21, 4, fill=1, stroke=0)
            c.setFillColor(colors.white); c.setFont("Helvetica-Bold", 7.6)
            c.drawCentredString(x + cw / 2, top - 14.5, label)
            body_h = top - 21 - 30
            c.setFillColor(colors.HexColor("#FCFBF6")); c.setStrokeColor(SAGE_LINE)
            c.roundRect(x, 24, cw, body_h, 4, fill=1, stroke=1)
            y = top - 21 - 18
            for n in names:
                rc = ROUTE_BG.get(route_of.get(n, ""), GREEN)
                c.setFillColor(rc); c.circle(x + 12, y + 3, 3, fill=1, stroke=0)
                c.setFillColor(INK); c.setFont("Helvetica-Bold", 8.2)
                c.drawString(x + 20, y, (n[:31] + "…") if len(n) > 32 else n)
                y -= 22
        c.setFont("Helvetica-Bold", 7.2); lx = 0
        for route, rc in [("Hotelería y spa", ROUTE_BG["Hotelería y spa"]), ("Retail especializado", GREEN),
                          ("Regalos corporativos", GOLD)]:
            c.setFillColor(rc); c.circle(lx + 4, 8, 3, fill=1, stroke=0)
            c.setFillColor(MUTED); c.drawString(lx + 11, 5, route); lx += len(route) * 4.6 + 34

class DocCanvas(canvasmod.Canvas):
    """Two-pass: cover unnumbered; body pages numbered 1..N with 'Página X de N'."""
    def __init__(self, *a, **k):
        super().__init__(*a, **k); self._saved = []
    def showPage(self):
        self._saved.append(dict(self.__dict__)); self._startPage()
    def save(self):
        n = len(self._saved)
        for state in self._saved:
            self.__dict__.update(state)
            if self._pageNumber > 1:
                self._chrome(n)
            super().showPage()
        super().save()
    def _chrome(self, n):
        self.saveState()
        self.setStrokeColor(SAGE_LINE); self.setLineWidth(0.6)
        self.line(LM, PAGE_H - 44, PAGE_W - RM, PAGE_H - 44)
        self.setFont("Helvetica-Bold", 7.3); self.setFillColor(NAVY)
        self.drawString(LM, PAGE_H - 40, "LEADLENS")
        self.setFont("Helvetica", 7.3); self.setFillColor(MUTED)
        self.drawString(LM + 46, PAGE_H - 40, "Amor de Gea · Piloto 1 · Portafolio de oportunidades B2B")
        self.setStrokeColor(SAGE_LINE); self.line(LM, 40, PAGE_W - RM, 40)
        self.setFont("Helvetica", 7.3); self.setFillColor(MUTED)
        self.drawString(LM, 30, "Documento para revisión y conversación comercial · Confidencial")
        self.drawRightString(PAGE_W - RM, 30, f"Página {self._pageNumber - 1} de {n - 1}")
        self.restoreState()

def cover(canvas, doc):
    canvas.saveState()
    canvas.setFillColor(NAVY); canvas.rect(0, 0, PAGE_W, PAGE_H, fill=1, stroke=0)
    canvas.setFillColor(colors.HexColor("#1E4034")); canvas.rect(0, PAGE_H - 250, PAGE_W, 250, fill=1, stroke=0)
    canvas.setStrokeColor(GOLD); canvas.setLineWidth(2.2); canvas.line(LM, PAGE_H - 150, LM + 70, PAGE_H - 150)
    canvas.setFillColor(GOLD_SOFT); canvas.setFont("Helvetica-Bold", 12); canvas.drawString(LM, PAGE_H - 128, "LEADLENS")
    canvas.setFillColor(colors.white); canvas.setFont("Helvetica", 10.5); canvas.drawString(LM, PAGE_H - 178, "INTELIGENCIA DE OPORTUNIDADES B2B")
    canvas.setFillColor(colors.white); canvas.setFont("Helvetica-Bold", 30)
    canvas.drawString(LM, PAGE_H - 320, "Portafolio de"); canvas.drawString(LM, PAGE_H - 356, "oportunidades B2B")
    canvas.setFillColor(GOLD_SOFT); canvas.drawString(LM, PAGE_H - 400, "y plan de validación"); canvas.drawString(LM, PAGE_H - 436, "comercial")
    canvas.setStrokeColor(GOLD); canvas.setLineWidth(1.2); canvas.line(LM, PAGE_H - 460, PAGE_W - RM, PAGE_H - 460)
    canvas.setFillColor(colors.white); canvas.setFont("Helvetica-Bold", 15); canvas.drawString(LM, PAGE_H - 492, "AMOR DE GEA")
    canvas.setFillColor(GOLD_SOFT); canvas.setFont("Helvetica", 10.5); canvas.drawString(LM, PAGE_H - 512, "Piloto LeadLens 1 · Infusiones botánicas · Colombia")
    canvas.setFillColor(colors.HexColor("#B9C9BE")); canvas.setFont("Helvetica", 10)
    canvas.drawString(LM, 150, "Diez cuentas · tres rutas comerciales · cuatro validaciones iniciales")
    canvas.setFillColor(colors.white); canvas.setFont("Helvetica-Bold", 10.5); canvas.drawString(LM, 120, DATA["meta"]["generated_label"])
    canvas.setFillColor(colors.HexColor("#8FA697")); canvas.setFont("Helvetica", 8.3)
    canvas.drawString(LM, 66, "Preparado por LeadLens para la conversación comercial de Amor de Gea.")
    canvas.drawString(LM, 52, "Las relaciones con cada cuenta deben confirmarse antes de cualquier contacto.")
    canvas.restoreState()

def _body_doc(out, title, cover_fn=None):
    frame = Frame(LM, 48, CONTENT_W, PAGE_H - 48 - 56, id="body")
    doc = BaseDocTemplate(str(out), pagesize=letter, title=title, author="LeadLens")
    if cover_fn:
        doc.addPageTemplates([PageTemplate(id="cover", frames=[frame], onPage=cover_fn),
                              PageTemplate(id="body", frames=[frame])])
    else:
        doc.addPageTemplates([PageTemplate(id="body", frames=[frame])])
    return doc

def brief_body(s, b):
    """One-page detailed brief body (shared by the report and the briefs package)."""
    callout(s, b["thesis"], after=6)
    table(s, ["Elemento", "Detalle"],
          [["Prueba recomendada", b["test"]], ["Función compradora (hipótesis)", b["buyer_hyp"]],
           ["Estructura de decisión", b["procurement"]], ["Ciclo comercial", b["cycle"]],
           ["Por validar antes de contactar", b["next"]]], [148, CONTENT_W - 148])
    s.append(Spacer(1, 5))
    ev = b["evidence"]
    table(s, [f"Fuente · consultado {ev['retrieved']}", "Hecho público", "No confirma"],
          [[ev["source"] + " · sitio oficial", ev["fact"], ev["not_proves"]]], [110, 196, 194])
    s.append(Spacer(1, 6))
    qcol = [P("<b>Preguntas clave</b>", "CardBody")] + [P("• " + q, "CardBody") for q in b["questions"]]
    ocol = [P("<b>Objeciones probables (hipótesis)</b>", "CardBody")] + [P("• " + o, "CardBody") for o in b["objections"]] + [P("<b>Materiales:</b> " + " · ".join(b["prep"]), "Small")]
    two = Table([[qcol, ocol]], colWidths=[CONTENT_W / 2, CONTENT_W / 2])
    two.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (0, 0), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 8)]))
    s.append(two); s.append(Spacer(1, 6))
    gcol = [P("<b>Guía de decisión</b>", "CardBody"),
            P("<b>Avanzar:</b> uso específico + responsable + economía plausible + siguiente paso", "CardBody"),
            P("<b>Pausar:</b> interés general sin decisión, datos o plazo", "CardBody"),
            P("<b>Descartar:</b> sin caso de uso, conflicto, margen inviable o formato incompatible", "CardBody")]
    ncol = [P("<b>Notas de la conversación</b>", "CardBody"),
            P("Caso de uso: " + "_" * 30, "CardBody"), P("Responsable / área: " + "_" * 24, "CardBody"),
            P("Economía y pedido mínimo: " + "_" * 18, "CardBody"), P("Siguiente paso y fecha: " + "_" * 20, "CardBody")]
    gt = Table([[gcol, ncol]], colWidths=[CONTENT_W * 0.52, CONTENT_W * 0.48])
    gt.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP"), ("BACKGROUND", (0, 0), (-1, -1), CREAM),
                            ("BOX", (0, 0), (-1, -1), 0.5, SAGE_LINE), ("LEFTPADDING", (0, 0), (-1, -1), 10),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 10), ("TOPPADDING", (0, 0), (-1, -1), 7), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
    s.append(gt)
    s.append(P("No interpretar interés cordial como intención de compra. No prometer efectos de salud, volumen ni fechas.", "Small"))

CONTENTS = [
    ("01", "Conclusión ejecutiva"), ("02", "Contexto y cómo cambió la búsqueda"),
    ("03", "Preparación comercial y rutas"), ("04", "Mapa de oportunidades y evidencia"),
    ("05", "Portafolio de un vistazo"), ("06", "Primera secuencia de validación"),
    ("07", "Prioridad estratégica"), ("08", "Investigación selectiva y por qué no ahora"),
    ("09", "Briefs de acción por cuenta"), ("10", "Qué cambió y preparación comercial"),
    ("11", "Plan 30–60 días y marco de éxito"), ("12", "Evidencia, límites y cierre"),
]

def build_report():
    out = PDF / "Amor-de-Gea-LeadLens-Pilot-1-Final-Report.pdf"
    doc = _body_doc(out, "Amor de Gea · Portafolio de oportunidades B2B y plan de validación comercial", cover)
    s = [NextPageTemplate("body"), PageBreak()]

    # 01 Conclusión + contenido
    section(s, "01", "Conclusión ejecutiva")
    callout(s, "Recomendamos comenzar con cuatro validaciones pequeñas y distintas: Éteka y Celestino en hotelería/spa, Sinergy On en regalos corporativos y Vitálica en retail natural.", after=7)
    s.append(P("El portafolio no afirma intención de compra: organiza dónde aprender primero, qué validar y por qué, con pruebas compatibles con la capacidad actual de Amor de Gea.", "BodyL"))
    bullets(s, ["Diez cuentas activas, organizadas por secuencia y nivel de validación.",
                "Se privilegian pruebas accesibles y mecanismos de recompra, no empresas famosas por su nombre.",
                "El éxito se mide por aprendizaje comercial verificable, no por promesas de venta."])
    s.append(Spacer(1, 4))
    s.append(P("<b>Decisión sugerida:</b> revisar la relación previa con cada cuenta, preparar cuatro conversaciones de descubrimiento y ejecutar solo pruebas comerciales acotadas.", "BodyL"))
    s.append(Spacer(1, 8)); s.append(P("CONTENIDO", "Kicker"))
    half = (len(CONTENTS) + 1) // 2
    nav_rows = [[f'{CONTENTS[i][0]} · {CONTENTS[i][1]}', f'{CONTENTS[i+half][0]} · {CONTENTS[i+half][1]}' if i + half < len(CONTENTS) else ""] for i in range(half)]
    table(s, ["Sección", "Sección"], nav_rows, [CONTENT_W / 2, CONTENT_W / 2])
    s.append(PageBreak())

    # 02 Contexto + cómo cambió la búsqueda
    section(s, "02", "Contexto y cómo cambió la búsqueda")
    s.append(P("<b>Qué entendió LeadLens de Amor de Gea:</b>", "BodyL"))
    bullets(s, ["Tres elixires botánicos terminados con presentación premium (Agua, Tierra, Éter).",
                "Escalamiento progresivo hacia ~300 unidades mensuales; pedido mínimo de piloto cercano a 50 unidades.",
                "Interés en retail especializado, hotelería/spa y regalos corporativos.",
                "Restricciones: sin lenguaje médico, validar documentación, manejo del vidrio, márgenes y personalización."])
    s.append(Spacer(1, 5))
    wc = DATA["what_changed"]
    n = max(len(wc["before"]), len(wc["after"]))
    rows = [[wc["before"][i] if i < len(wc["before"]) else "", wc["after"][i] if i < len(wc["after"]) else ""] for i in range(n)]
    table(s, ["Antes del contexto completo", "Después del contexto completo"], rows, [CONTENT_W / 2, CONTENT_W / 2])
    s.append(PageBreak())

    # 03 Preparación comercial + rutas
    section(s, "03", "Preparación comercial y rutas")
    rd = DATA["readiness"]
    table(s, ["Fortalezas hoy", "Validar antes de vender"],
          [[st, va] for st, va in zip(rd["strengths"], rd["validate"])], [CONTENT_W / 2, CONTENT_W / 2])
    s.append(Spacer(1, 6))
    bullets(s, [f'<b>{m["route"].title()}:</b> {m["summary"]}' for m in DATA["market_map"]], "CardBody")
    s.append(Spacer(1, 4))
    s.append(P("Recomendación: no escalar hasta observar una señal real de uso, economía y posibilidad de reposición.", "Small"))
    s.append(PageBreak())

    # 04 Mapa de oportunidades + evidencia por ruta
    section(s, "04", "Mapa de oportunidades y evidencia por ruta")
    table(s, ["Ruta", "Cuenta de entrada", "Mecanismo a validar"],
          [["Hotelería / spa", "Éteka · Celestino", "Ritual, regalo o venta complementaria"],
           ["Regalos corporativos", "Sinergy On", "Producto terminado dentro de un kit"],
           ["Retail natural", "Vitálica", "Surtido pequeño y rotación de venta"],
           ["Expansión estratégica", "Ser Saludable · Masaya · Natural + Mente", "Replicar solo tras el aprendizaje inicial"]],
          [108, 165, 207])
    s.append(Spacer(1, 6))
    table(s, ["Ruta", "Evidencia", "Mecanismo", "Conclusión"],
          [[r["route"], r["evidence"], r["mechanism"], r["conclusion"]] for r in DATA["route_review"]],
          [92, 88, 150, 150])
    s.append(PageBreak())

    # 05 Portafolio de un vistazo
    section(s, "05", "Portafolio de un vistazo")
    s.append(PortfolioMap(DATA)); s.append(Spacer(1, 8))
    s.append(P("Las tres columnas indican secuencia y nivel de validación; el color marca la ruta de cada cuenta.", "Small"))
    s.append(P(DATA["relationship_disclosure"], "Small"))
    s.append(PageBreak())

    # 06 Primera secuencia (full evidence cards)
    section(s, "06", "Primera secuencia de validación")
    s.append(P("Cuatro cuentas con evidencia pública actual y una prueba inicial compatible. Cada tarjeta muestra la fuente, el hecho público, lo que confirma y lo que no.", "BodyL"))
    for a in [x for x in DATA["accounts"] if x["group"] == "Primera validación"]:
        s.append(account_card(a, full_evidence=True))
    s.append(PageBreak())

    # 07 Prioridad estratégica (compact cards)
    section(s, "07", "Prioridad estratégica")
    s.append(P("Cuentas relevantes para un segundo momento: encaje estructural probable, con validación adicional antes de priorizarlas.", "BodyL"))
    for a in [x for x in DATA["accounts"] if x["group"] == "Prioridad estratégica"]:
        s.append(account_card(a, full_evidence=False))
    s.append(PageBreak())

    # 08 Investigación selectiva + por qué no ahora
    section(s, "08", "Investigación selectiva y por qué no ahora")
    for a in [x for x in DATA["accounts"] if x["group"] == "Investigar selectivamente"]:
        s.append(account_card(a, full_evidence=False))
    s.append(Spacer(1, 4))
    s.append(KeepTogether([
        P("Cuentas no priorizadas ahora", "H3x"),
        P("No es un rechazo definitivo: se conservan como aprendizaje y pueden reconsiderarse con nueva evidencia.", "Small"),
        mk_table(["Cuenta", "Por qué no ahora"], [[e["name"], e["reason"]] for e in DATA["excluded"]], [130, CONTENT_W - 130])]))
    s.append(PageBreak())

    # 09 Briefs (one detailed page each)
    for i, b in enumerate(DATA["briefs"], 1):
        section(s, "09", b["name"], kicker=f"Sección 09 · Brief de acción {i} de 4 · {b['route']}")
        brief_body(s, b)
        s.append(PageBreak())

    # 10 Qué cambió + preparación
    section(s, "10", "Qué cambió y preparación comercial")
    bullets(s, DATA["what_changed"]["narrative"], "CardBody")
    s.append(Spacer(1, 5)); s.append(P("Lista de preparación comercial", "H3x"))
    bullets(s, DATA["prep_checklist"], "CardBody")
    s.append(PageBreak())

    # 11 Plan + marco de éxito
    section(s, "11", "Plan de validación y marco de éxito")
    table(s, ["Periodo", "Acción", "Resultado esperado"],
          [["Días 1–10", "Revisión de relaciones, conflicto y materiales", "Cuatro cuentas autorizadas para descubrimiento"],
           ["Días 11–25", "Conversaciones de validación", "Necesidad, uso, comprador, economía y objeciones"],
           ["Días 26–40", "Uno o dos conceptos de prueba", "Tamaño, responsable, medición y salida"],
           ["Días 41–60", "Decisión de continuar o descartar", "Aprendizaje por ruta y foco del siguiente ciclo"]],
          [78, 190, 212])
    s.append(Spacer(1, 6)); s.append(P(f'<b>Objetivo del piloto:</b> {DATA["success"]["objective"]}', "BodyL"))
    bullets(s, DATA["success"]["value"], "CardBody")
    s.append(Spacer(1, 4)); callout(s, DATA["success"]["no_guarantee"], after=2)
    s.append(PageBreak())

    # 12 Evidencia + cierre
    section(s, "12", "Evidencia, límites y cierre")
    bullets(s, DATA["limitations"], "CardBody")
    s.append(Spacer(1, 6)); callout(s, "Piloto 1 completado — portafolio preparado para validación comercial.", after=6)
    s.append(P(DATA["closing_pilot2"], "BodyL"))
    s.append(P("Antes de cualquier contacto, confirme la relación previa con cada cuenta.", "Small"))

    doc.build(s, canvasmaker=DocCanvas)
    return out

def build_briefs():
    out = PDF / "Amor-de-Gea-Account-Action-Briefs-Pilot-1.pdf"
    doc = _body_doc(out, "Amor de Gea · Briefs de acción por cuenta")
    s = [Spacer(1, 26), P("LEADLENS · AMOR DE GEA", "Kicker"), P("Briefs de acción por cuenta", "H1x"),
         Rule(width=64, color=GOLD, thick=2.4, gap=5), Spacer(1, 7),
         P("Cuatro cuentas de la primera secuencia de validación. Cada brief se apoya en evidencia pública atribuida y en el contexto de Amor de Gea. No incluye contactos personales ni intención de compra: describe hipótesis a validar antes de cualquier acuerdo comercial. Antes de contactar cualquier cuenta, confirme la relación previa o un posible conflicto.", "BodyL")]
    s.append(Spacer(1, 6))
    s.append(P("Cómo usar cada brief", "H3x"))
    bullets(s, ["Lleve muestra, ficha, condiciones y lenguaje seguro (no médico).",
                "Use las preguntas específicas de la cuenta para dirigir la conversación.",
                "Registre caso de uso, responsable, economía, objeciones y siguiente paso.",
                "Aplique la guía de decisión (avanzar / pausar / descartar) al cierre."], "CardBody")
    s.append(PageBreak())
    for i, b in enumerate(DATA["briefs"], 1):
        s.append(P(f"BRIEF {i} DE 4 · {b['route'].upper()}", "Kicker"))
        s.append(P(b["name"], "H1x")); s.append(Rule(width=64, color=GOLD, thick=2.4, gap=5)); s.append(Spacer(1, 7))
        brief_body(s, b)
        if i < len(DATA["briefs"]):
            s.append(PageBreak())
    doc.build(s, canvasmaker=DocCanvas)
    return out

# ---- Feedback --------------------------------------------------------------
ACC = [a["name"] for a in DATA["accounts"]]
FIRST4 = [a["name"] for a in DATA["accounts"] if a["group"] == "Primera validación"]
G = DATA["feedback_guide"]

def _writelines(s, k=1):
    for _ in range(k):
        s.append(P("_" * 92, "Small"))

def build_feedback_pdf():
    out = PDF / "Amor-de-Gea-LeadLens-Pilot-1-Feedback.pdf"
    doc = _body_doc(out, "Amor de Gea · Retroalimentación Piloto 1")
    s = [Spacer(1, 26), P("LEADLENS · AMOR DE GEA", "Kicker"), P("Retroalimentación · Piloto 1", "H1x"),
         Rule(width=64, color=GOLD, thick=2.4, gap=5), Spacer(1, 6),
         P("Sus respuestas convertirán este portafolio en aprendizaje para el siguiente ciclo. No hay respuestas prellenadas: marque las casillas y escriba en los espacios provistos.", "BodyL")]
    # Cómo evaluar este piloto
    s.append(P("Cómo evaluar este piloto", "H2x"))
    s.append(P(G["how_to_evaluate"]["intro"], "CardBody"))
    table(s, ["Dimensión", "Qué observar"],
          [[f'{d["n"]}. {d["title"]}', d["q"]] for d in G["how_to_evaluate"]["dimensions"]], [130, CONTENT_W - 130])
    s.append(Spacer(1, 4)); callout(s, G["how_to_evaluate"]["note"], after=4)
    s.append(PageBreak())

    # Escala + Quién responde
    section(s, "A", "Cómo usar la escala y quién responde", kicker="Sección A")
    s.append(P("Escala 1–5", "H3x"))
    for lv in G["scale_guide"]["levels"]:
        s.append(P(f'<b>{lv["n"]}</b> — {lv["text"]}', "CardBody"))
    bullets(s, G["scale_guide"]["guidance"], "Small")
    s.append(Spacer(1, 6)); s.append(P("Quién responde", "H3x"))
    for q in ["Nombre", "Cargo o relación con Amor de Gea", "Fecha"]:
        s.append(P(f"<b>{q}</b>", "CardBody")); _writelines(s, 1)
    s.append(PageBreak())

    # B general ratings
    section(s, "B", "Valoración general del piloto", kicker="Sección B")
    s.append(P("Califique 1–5 cada aspecto (vea la escala en la sección A).", "Small"))
    for q in ["Utilidad general", "Relevancia de las cuentas", "Utilidad de la priorización",
              "Utilidad de los briefs de acción", "Credibilidad de la evidencia", "Uso del contexto",
              "Claridad del reporte", "Confianza para decidir próximos pasos",
              "Probabilidad de usar un Piloto 2", "Valor frente a una base de datos"]:
        s.append(P(f"<b>{q}</b>", "CardBody")); s.append(checkbox_row(["1", "2", "3", "4", "5"]))
    s.append(Spacer(1, 8)); s.append(P("<b>Comentario general sobre el valor del piloto</b>", "CardBody"))
    _writelines(s, 3)
    s.append(PageBreak())

    # Account guide + C (2 pages)
    section(s, "C", "Retroalimentación de las 10 cuentas", kicker="Sección C")
    s.append(P("Para cada cuenta, considere:", "CardBody"))
    bullets(s, G["account_guide"]["questions"], "Small")
    s.append(P("Estado de la relación: " + "  ".join(f'<b>{r["label"]}</b> ({r["desc"]})' for r in G["account_guide"]["relationship"]), "Small"))
    s.append(Spacer(1, 4))
    for idx, name in enumerate(ACC):
        block = [P(f"<b>{name}</b>", "CardBody"),
                 checkbox_row(["Nueva", "Conocida", "Contactada", "Conversación activa", "Cliente/socio", "Excluir"]),
                 checkbox_row(["Relevancia alta", "media", "baja", "  ·  Validar primero", "Mantener", "Investigar", "Descartar"]),
                 P("Nota: " + "_" * 78, "Small")]
        s.append(KeepTogether(block)); s.append(Spacer(1, 3))
        if idx == 4:
            s.append(PageBreak())
    s.append(PageBreak())

    # D first sequence + brief guide + E
    section(s, "D", "Primera secuencia y evaluación de los briefs", kicker="Secciones D–E")
    s.append(P("<b>¿Qué cuentas validaría primero?</b> Marque las que priorizaría:", "CardBody"))
    s.append(checkbox_row(FIRST4))
    s.append(P("Indique cuentas conocidas, contactadas, activas o que deban excluirse: " + "_" * 34, "Small"))
    s.append(Spacer(1, 6)); s.append(P("Evaluación de los cuatro briefs de acción", "H3x"))
    s.append(P("Para cada brief, considere:", "CardBody"))
    bullets(s, G["brief_guide"], "Small")
    for name in FIRST4:
        s.append(P(f"<b>{name}</b>", "CardBody")); s.append(checkbox_row(["1", "2", "3", "4", "5"]))
        s.append(P("Comentario: " + "_" * 74, "Small"))
    s.append(PageBreak())

    # F routes + G readiness
    section(s, "F", "Preferencias por ruta y correcciones", kicker="Secciones F–G")
    s.append(P("Preferencia por ruta (1–5)", "H3x"))
    for r in ["Hotelería / spa", "Retail natural", "Regalos corporativos", "Distribución u otra"]:
        s.append(P(f"<b>{r}</b>", "CardBody")); s.append(checkbox_row(["1", "2", "3", "4", "5"]))
    s.append(Spacer(1, 5)); s.append(P("Correcciones de preparación comercial", "H3x"))
    for q in ["Precio y margen", "Capacidad y pedido mínimo", "Documentación", "Personalización", "Logística y vidrio"]:
        s.append(P(f"<b>{q}</b>: " + "_" * 60, "CardBody"))
    s.append(PageBreak())

    # Novelty + decision + H + pilot 2
    section(s, "H", "Valor del piloto y prioridades del Piloto 2", kicker="Secciones H–I")
    for q in G["key_questions"]:
        s.append(P(f"<b>{q}</b>", "CardBody")); _writelines(s, 2)
    for q in ["¿Qué decisión permitió tomar el piloto?", "¿Qué faltó para generar más valor?"]:
        s.append(P(f"<b>{q}</b>", "CardBody")); _writelines(s, 1)
    s.append(Spacer(1, 4)); s.append(P("Prioridades del Piloto 2", "H3x"))
    s.append(P("Un segundo ciclo no debería repetir las mismas cuentas. Indique:", "CardBody"))
    for item in G["pilot2_guide"]:
        s.append(P(f"• {item} " + "_" * 22, "Small"))
    s.append(PageBreak())

    # J commercial + K comments/consent
    section(s, "J", "Disposición comercial y comentarios", kicker="Secciones J–K")
    for q in G["commercial"]["questions"]:
        s.append(P(f"<b>{q}</b>", "CardBody")); _writelines(s, 1)
    s.append(P("<b>Formato más valioso</b> (marque una o varias):", "CardBody"))
    s.append(checkbox_row(G["commercial"]["formats"]))
    s.append(Spacer(1, 6)); s.append(P("Comentarios y consentimiento", "H3x"))
    s.append(P("<b>Comentarios finales</b>", "CardBody")); _writelines(s, 2)
    s.append(P("<b>¿Autoriza usar esta respuesta para mejorar LeadLens?</b>", "CardBody"))
    s.append(checkbox_row(["Sí", "No"]))
    doc.build(s, canvasmaker=DocCanvas)
    return out

def _shade(cell, fill):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); cell._tc.get_or_add_tcPr().append(shd)

def _answer_cell(d, height_lines=2):
    tb = d.add_table(rows=1, cols=1); tb.autofit = False; tb.alignment = WD_TABLE_ALIGNMENT.LEFT
    tb.columns[0].width = Inches(6.5); c = tb.cell(0, 0); c.width = Inches(6.5)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP; c.text = "Respuesta / nota:" + "\n" * height_lines
    _shade(c, "F8F5ED")

def build_feedback_docx():
    out = DOCX / "Amor-de-Gea-LeadLens-Pilot-1-Feedback.docx"
    d = Document(); sec = d.sections[0]
    sec.page_width = Inches(8.5); sec.page_height = Inches(11)
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    sec.header_distance = sec.footer_distance = Inches(0.5)
    normal = d.styles["Normal"]; normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(5); normal.paragraph_format.line_spacing = 1.1
    for sty, size, color, before, after in [("Heading 1", 15, "17352C", 14, 5), ("Heading 2", 12.5, "4E6A54", 11, 4), ("Heading 3", 11, "6B7873", 8, 3)]:
        x = d.styles[sty]; x.font.name = "Calibri"; x.font.size = Pt(size); x.font.color.rgb = RGBColor.from_string(color)
        x.paragraph_format.space_before = Pt(before); x.paragraph_format.space_after = Pt(after); x.paragraph_format.keep_with_next = True
    hp = sec.header.paragraphs[0]; hp.text = "LEADLENS  ·  AMOR DE GEA · PILOTO 1"; hp.runs[0].font.size = Pt(8); hp.runs[0].font.color.rgb = RGBColor.from_string("6B7873")
    fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("Documento de retroalimentación · v1.2 · las respuestas se revisan antes de usarse"); fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string("6B7873")
    k = d.add_paragraph(); k.paragraph_format.space_before = Pt(24); r = k.add_run("LEADLENS · AMOR DE GEA"); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string("B48A4A")
    t = d.add_paragraph(); r = t.add_run("Retroalimentación · Piloto LeadLens 1"); r.bold = True; r.font.size = Pt(23); r.font.color.rgb = RGBColor.from_string("17352C")
    d.add_paragraph("Sus respuestas convertirán este portafolio en aprendizaje para el siguiente ciclo. No hay respuestas prellenadas. Puede escribir directamente en este documento.")

    d.add_heading("Cómo evaluar este piloto", level=2)
    d.add_paragraph(G["how_to_evaluate"]["intro"])
    for dim in G["how_to_evaluate"]["dimensions"]:
        p = d.add_paragraph(style="List Bullet"); p.add_run(f'{dim["n"]}. {dim["title"]}: ').bold = True; p.add_run(dim["q"])
    d.add_paragraph(G["how_to_evaluate"]["note"])
    d.add_paragraph()
    d.add_heading("Cómo usar la escala 1–5", level=3)
    for lv in G["scale_guide"]["levels"]:
        p = d.add_paragraph(); p.add_run(f'{lv["n"]} — ').bold = True; p.add_run(lv["text"])
    for gtext in G["scale_guide"]["guidance"]:
        d.add_paragraph(gtext, style="List Bullet")
    d.add_page_break()

    def rating(label):
        p = d.add_paragraph(); p.add_run(label).bold = True
        d.add_paragraph("Calificación 1–5:   ☐ 1    ☐ 2    ☐ 3    ☐ 4    ☐ 5")

    d.add_heading("A. Quién responde", level=1)
    for q in ["Nombre", "Cargo o relación con Amor de Gea", "Fecha"]:
        d.add_paragraph().add_run(q).bold = True; _answer_cell(d, 1)
    d.add_heading("B. Valoración general del piloto", level=1)
    for q in ["Utilidad general", "Relevancia de las cuentas", "Utilidad de la priorización",
              "Utilidad de los briefs de acción", "Credibilidad de la evidencia", "Uso del contexto",
              "Claridad del reporte", "Confianza para decidir próximos pasos",
              "Probabilidad de usar un Piloto 2", "Valor frente a una base de datos"]:
        rating(q)
    d.add_page_break()

    d.add_heading("C. Retroalimentación de las 10 cuentas", level=1)
    d.add_paragraph("Para cada cuenta, considere:")
    for q in G["account_guide"]["questions"]:
        d.add_paragraph(q, style="List Bullet")
    d.add_paragraph("Estado de la relación: " + "; ".join(f'{r["label"]} ({r["desc"]})' for r in G["account_guide"]["relationship"]))
    for name in ACC:
        d.add_paragraph().add_run(name).bold = True
        d.add_paragraph("Relación:  ☐ Nueva  ☐ Conocida  ☐ Contactada  ☐ Conversación activa  ☐ Cliente/socio  ☐ Excluir")
        d.add_paragraph("Relevancia:  ☐ alta  ☐ media  ☐ baja      Acción:  ☐ validar primero  ☐ mantener  ☐ investigar  ☐ descartar")
        _answer_cell(d, 1)
    d.add_page_break()

    d.add_heading("D. Primera secuencia", level=1)
    d.add_paragraph("¿Qué cuentas validaría primero? " + "   ".join(f"☐ {n}" for n in FIRST4))
    d.add_paragraph().add_run("Cuentas conocidas, contactadas, activas o a excluir:").bold = True; _answer_cell(d, 1)
    d.add_heading("E. Evaluación de los cuatro briefs de acción", level=1)
    d.add_paragraph("Para cada brief, considere:")
    for q in G["brief_guide"]:
        d.add_paragraph(q, style="List Bullet")
    for name in FIRST4:
        rating(name); _answer_cell(d, 1)
    d.add_page_break()

    d.add_heading("F. Preferencias por ruta", level=1)
    for r in ["Hotelería / spa", "Retail natural", "Regalos corporativos", "Distribución u otra"]:
        rating(r)
    d.add_heading("G. Correcciones de preparación comercial", level=1)
    for q in ["Precio y margen", "Capacidad y pedido mínimo", "Documentación", "Personalización", "Logística y vidrio"]:
        d.add_paragraph().add_run(q).bold = True; _answer_cell(d, 1)
    d.add_page_break()

    d.add_heading("H. Valor del piloto", level=1)
    for q in G["key_questions"] + ["¿Qué decisión permitió tomar el piloto?", "¿Qué faltó para generar más valor?"]:
        d.add_paragraph().add_run(q).bold = True; _answer_cell(d, 2)
    d.add_heading("I. Prioridades del Piloto 2", level=1)
    d.add_paragraph("Un segundo ciclo no debería repetir las mismas cuentas. Indique:")
    for item in G["pilot2_guide"]:
        d.add_paragraph(item, style="List Bullet")
    _answer_cell(d, 2)
    d.add_page_break()

    d.add_heading("J. Disposición comercial", level=1)
    for q in G["commercial"]["questions"]:
        d.add_paragraph().add_run(q).bold = True; _answer_cell(d, 1)
    d.add_paragraph("Formato más valioso: " + "   ".join(f"☐ {f}" for f in G["commercial"]["formats"]))
    d.add_heading("K. Comentarios y consentimiento", level=1)
    d.add_paragraph().add_run("Comentarios finales").bold = True; _answer_cell(d, 3)
    d.add_paragraph("¿Autoriza usar esta respuesta para mejorar LeadLens?   ☐ Sí    ☐ No")

    cp = d.core_properties
    cp.author = "LeadLens"; cp.title = "Amor de Gea · Retroalimentación Piloto 1"
    cp.created = cp.modified = datetime(2026, 8, 3); cp.last_modified_by = "LeadLens"; cp.revision = 1
    d.save(out); _normalize_zip(out)
    return out

def _normalize_zip(path):
    import zipfile
    with zipfile.ZipFile(path) as z:
        data = {i.filename: z.read(i.filename) for i in z.infolist()}
    tmp = path.with_suffix(".docx.tmp")
    with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as z:
        for name in sorted(data):
            zi = zipfile.ZipInfo(name, date_time=(2026, 8, 3, 0, 0, 0)); zi.compress_type = zipfile.ZIP_DEFLATED
            zi.external_attr = 0o644 << 16; z.writestr(zi, data[name])
    tmp.replace(path)

DENY = ["V3R3", "V3R2", "V4D", "Blueprint", "compiler", "provider", "Founder Review",
        "Revisión del fundador", "revisión interna", "No enviado", "admin_entry",
        "Phase 5", "conflict check", "NEEDS EVIDENCE", "actionability", "sell-through",
        "co-branding", "onboarding", "gifting", "MOQ"]

def _selfcheck():
    from pypdf import PdfReader
    import re
    report = "\n".join((p.extract_text() or "") for p in PdfReader(str(PDF / "Amor-de-Gea-LeadLens-Pilot-1-Final-Report.pdf")).pages)
    briefs = "\n".join((p.extract_text() or "") for p in PdfReader(str(PDF / "Amor-de-Gea-Account-Action-Briefs-Pilot-1.pdf")).pages)
    feedback = "\n".join((p.extract_text() or "") for p in PdfReader(str(PDF / "Amor-de-Gea-LeadLens-Pilot-1-Feedback.pdf")).pages)
    for label, text in [("report", report), ("briefs", briefs), ("feedback", feedback)]:
        for tok in DENY:
            assert not re.search(re.escape(tok), text, re.I), f"internal/English token '{tok}' leaked into {label}"
    for n in ACC:
        assert n in report, f"missing account {n} in report"
        assert n in feedback, f"missing account {n} in feedback"
    for src in ["etekacartagena.com", "hotelcelestino.com", "sinergyon.com", "tiendavitalica.com"]:
        assert src in report and src in briefs, f"missing evidence source {src}"
    assert "Piloto 1 completado" in report
    assert "Cómo evaluar este piloto" in feedback and "Cómo usar la escala" in feedback
    assert "genuinamente nueva" in feedback  # novelty question
    print("selfcheck: ok (no internal/English tokens · 10 accounts · evidence · feedback guides)")

def main():
    files = [build_report(), build_briefs(), build_feedback_pdf(), build_feedback_docx()]
    for f in files:
        shutil.copy2(f, PUBLIC / f.name)
    _selfcheck()
    from pypdf import PdfReader
    for f in files:
        pages = len(PdfReader(str(f)).pages) if f.suffix == ".pdf" else "-"
        print(f"{f.name}  ({f.stat().st_size} bytes · {pages} pages)")

if __name__ == "__main__":
    main()
